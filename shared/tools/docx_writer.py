from __future__ import annotations

from pathlib import Path
from typing import Iterable

from shared.schemas.template import TemplateFormattingProfile
from shared.tools.docx_renderer import apply_cjk_formatting


def new_document():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('DOCX 生成需要 python-docx：pip install -e ".[dev]"') from exc
    return Document()


def add_bullets(document, heading: str, items: Iterable[str]) -> None:
    document.add_heading(heading, level=2)
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def save_document(document, output_path: Path, formatting: TemplateFormattingProfile | None = None) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    apply_cjk_formatting(document, formatting or TemplateFormattingProfile())
    document.save(output_path)
