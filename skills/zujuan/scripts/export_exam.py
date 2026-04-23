from __future__ import annotations

import argparse
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from pydantic import ValidationError

from shared.schemas.exam import ExamPaper
from shared.schemas.question import Question


def load_exam(raw_json: str) -> ExamPaper:
    return ExamPaper.model_validate_json(raw_json)


def export_markdown(exam: ExamPaper, include_answers: bool = False) -> str:
    lines: list[str] = [
        f"# {exam.title}",
        "",
        f"- 学科：{exam.subject}",
        f"- 年级：{exam.grade}",
        f"- 满分：{exam.total_score:g}",
        f"- 时长：{exam.duration_minutes} 分钟",
        "",
    ]

    question_number = 1
    answer_lines: list[str] = ["", "## 参考答案", ""]
    for section in exam.sections:
        lines.extend([f"## {section.title}", "", f"本部分共 {len(section.questions)} 题，{section.section_score:g} 分。", ""])
        for question in section.questions:
            lines.extend(_render_question(question_number, question))
            if include_answers:
                answer_lines.extend(_render_answer(question_number, question))
            question_number += 1

    if include_answers:
        lines.extend(answer_lines)
    return "\n".join(lines).rstrip() + "\n"


def export_docx(exam: ExamPaper, output_path: Path, include_answers: bool = False) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('DOCX 导出需要安装 python-docx：pip install -e ".[export]"') from exc

    document = Document()
    document.add_heading(exam.title, level=1)
    document.add_paragraph(f"学科：{exam.subject}    年级：{exam.grade}")
    document.add_paragraph(f"满分：{exam.total_score:g}    时长：{exam.duration_minutes} 分钟")

    question_number = 1
    answers: list[tuple[int, Question]] = []
    for section in exam.sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(f"本部分共 {len(section.questions)} 题，{section.section_score:g} 分。")
        for question in section.questions:
            if question.material:
                document.add_paragraph(question.material)
            document.add_paragraph(f"{question_number}. ({question.score:g} 分) {question.content}")
            if question.options:
                for option in question.options:
                    document.add_paragraph(option)
            answers.append((question_number, question))
            question_number += 1

    if include_answers:
        document.add_page_break()
        document.add_heading("参考答案", level=1)
        for number, question in answers:
            document.add_paragraph(f"{number}. {question.answer}")
            if question.explanation:
                document.add_paragraph(f"解析：{question.explanation}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _render_question(number: int, question: Question) -> list[str]:
    lines: list[str] = []
    if question.material:
        lines.extend([f"> {question.material}", ""])
    lines.append(f"{number}. ({question.score:g} 分) {question.content}")
    if question.options:
        lines.extend(question.options)
    lines.append("")
    return lines


def _render_answer(number: int, question: Question) -> list[str]:
    lines = [f"{number}. {question.answer}"]
    if question.explanation:
        lines.append(f"   解析：{question.explanation}")
    lines.append("")
    return lines


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="导出 ExamPaper 为 Markdown 或 DOCX。")
    parser.add_argument("--input", required=True, type=Path, help="ExamPaper JSON 文件")
    parser.add_argument("--format", choices=["markdown", "docx"], required=True, help="导出格式")
    parser.add_argument("--output", required=True, type=Path, help="输出文件路径")
    parser.add_argument("--include-answers", action="store_true", help="是否附带参考答案")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    try:
        exam = load_exam(args.input.read_text(encoding="utf-8"))
        if args.format == "markdown":
            args.output.parent.mkdir(parents=True, exist_ok=True)
            args.output.write_text(export_markdown(exam, include_answers=args.include_answers), encoding="utf-8")
        else:
            export_docx(exam, args.output, include_answers=args.include_answers)
    except FileNotFoundError as exc:
        sys.stderr.write(f"输入文件不存在: {exc.filename}\n")
        return 1
    except (RuntimeError, ValidationError) as exc:
        sys.stderr.write(f"{exc}\n")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
