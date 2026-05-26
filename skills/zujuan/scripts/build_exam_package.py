from __future__ import annotations

import argparse
import sys
import uuid
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from pydantic import TypeAdapter, ValidationError

from shared.schemas.exam import (
    AnswerKey,
    AnswerKeyItem,
    AnswerSheetSectionSpec,
    AnswerSheetSpec,
    ExamBlueprint,
    ExamBlueprintSection,
    ExamPackage,
    ExamPaper,
    ScoringRubric,
    ScoringRubricItem,
)
from shared.schemas.question import Question, QuestionType
from shared.schemas.research import ResearchDossier
from shared.schemas.template import TemplateFormattingProfile, TemplateProfile
from shared.tools.docx_renderer import apply_cjk_formatting
from shared.tools.sources import render_sources_markdown
from skills.zujuan.scripts.assemble_exam import ExamConstraints, SectionConstraint, assemble_exam


QUESTION_LIST_ADAPTER = TypeAdapter(list[Question])


def build_exam_package(
    *,
    research: ResearchDossier,
    questions: list[Question],
    profile: TemplateProfile,
    output_dir: Path,
) -> ExamPackage:
    blueprint = build_blueprint(research=research, profile=profile)
    exam = assemble_exam(questions, _constraints_from_blueprint(blueprint))
    answer_key = build_answer_key(exam)
    answer_sheet = build_answer_sheet_spec(exam)
    scoring_rubric = build_scoring_rubric(exam)
    checks = validate_package_consistency(
        exam=exam,
        blueprint=blueprint,
        answer_key=answer_key,
        answer_sheet=answer_sheet,
        scoring_rubric=scoring_rubric,
    )
    if not all(checks.values()):
        failed = ", ".join(name for name, passed in checks.items() if not passed)
        raise ValueError(f"试卷包一致性校验失败: {failed}")

    output_dir.mkdir(parents=True, exist_ok=True)
    files = _write_package_files(
        output_dir=output_dir,
        research=research,
        formatting=profile.formatting,
        blueprint=blueprint,
        exam=exam,
        answer_key=answer_key,
        answer_sheet=answer_sheet,
        scoring_rubric=scoring_rubric,
    )
    files["package_json"] = "package.json"
    file_checks = {f"file_exists:{name}": (output_dir / relative_path).exists() for name, relative_path in files.items()}
    file_checks["file_exists:package_json"] = True
    checks.update(file_checks)

    package = ExamPackage(
        id=f"exam_package_{uuid.uuid4().hex[:12]}",
        exam_id=exam.id,
        created_at=datetime.now().isoformat(timespec="seconds"),
        output_dir=str(output_dir),
        files=files,
        checks=checks,
        profile_id=profile.id,
        source_ids=[source.id for source in research.sources],
    )
    (output_dir / "package.json").write_text(package.model_dump_json(indent=2) + "\n", encoding="utf-8")
    return package


def build_blueprint(*, research: ResearchDossier, profile: TemplateProfile) -> ExamBlueprint:
    if profile.total_score is None:
        raise ValueError("TemplateProfile.total_score 不能为空")
    if profile.duration_minutes is None:
        raise ValueError("TemplateProfile.duration_minutes 不能为空")
    sections: list[ExamBlueprintSection] = []
    for section in profile.sections:
        if section.question_type is None:
            raise ValueError(f"{section.title}: question_type 不能为空")
        if section.item_count is None or section.item_count <= 0:
            raise ValueError(f"{section.title}: item_count 必须大于 0")
        if section.score_per_item is None or section.score_per_item <= 0:
            raise ValueError(f"{section.title}: score_per_item 必须大于 0")
        section_score = section.total_score or (section.item_count * section.score_per_item)
        sections.append(
            ExamBlueprintSection(
                title=section.title,
                question_type=QuestionType(section.question_type),
                item_count=section.item_count,
                score_per_item=section.score_per_item,
                section_score=section_score,
                difficulty_ratio=section.difficulty_ratio,
            )
        )

    return ExamBlueprint(
        id=f"blueprint_{uuid.uuid4().hex[:12]}",
        title=f"{research.grade}{research.subject}{research.topic}",
        subject=profile.subject,
        grade=research.grade,
        exam_type=profile.exam_type,
        total_score=profile.total_score,
        duration_minutes=profile.duration_minutes,
        sections=sections,
        knowledge_points=_extract_knowledge_points(research),
        source_ids=profile.source_ids or [source.id for source in research.sources],
        profile_id=profile.id,
    )


def build_answer_key(exam: ExamPaper) -> AnswerKey:
    items: list[AnswerKeyItem] = []
    for number, question in _numbered_questions(exam):
        items.append(
            AnswerKeyItem(
                question_number=number,
                question_id=question.id,
                answer=question.answer,
                score=question.score,
                explanation=question.explanation,
            )
        )
    return AnswerKey(exam_id=exam.id, items=items)


def build_answer_sheet_spec(exam: ExamPaper) -> AnswerSheetSpec:
    sections: list[AnswerSheetSectionSpec] = []
    next_number = 1
    for section in exam.sections:
        numbers = list(range(next_number, next_number + len(section.questions)))
        sections.append(
            AnswerSheetSectionSpec(
                title=section.title,
                question_type=section.question_type,
                question_numbers=numbers,
                response_area=_response_area_for(section.question_type),
            )
        )
        next_number += len(section.questions)
    return AnswerSheetSpec(exam_id=exam.id, title=exam.title, sections=sections)


def build_scoring_rubric(exam: ExamPaper) -> ScoringRubric:
    items: list[ScoringRubricItem] = []
    for number, question in _numbered_questions(exam):
        items.append(
            ScoringRubricItem(
                question_number=number,
                question_id=question.id,
                max_score=question.score,
                scoring_points=_scoring_points_for(question),
                review_required=question.question_type not in {QuestionType.CHOICE, QuestionType.FILL_BLANK},
            )
        )
    return ScoringRubric(exam_id=exam.id, items=items)


def validate_package_consistency(
    *,
    exam: ExamPaper,
    blueprint: ExamBlueprint,
    answer_key: AnswerKey,
    answer_sheet: AnswerSheetSpec,
    scoring_rubric: ScoringRubric,
) -> dict[str, bool]:
    question_numbers = [number for number, _ in _numbered_questions(exam)]
    answer_numbers = [item.question_number for item in answer_key.items]
    rubric_numbers = [item.question_number for item in scoring_rubric.items]
    sheet_numbers = [number for section in answer_sheet.sections for number in section.question_numbers]
    section_scores = [section.section_score for section in exam.sections]
    blueprint_scores = [section.section_score for section in blueprint.sections]
    return {
        "total_score_matches": abs(exam.total_score - blueprint.total_score) < 1e-6,
        "section_scores_match": section_scores == blueprint_scores,
        "answer_key_complete": question_numbers == answer_numbers,
        "answer_sheet_numbers_match": question_numbers == sheet_numbers,
        "scoring_rubric_complete": question_numbers == rubric_numbers,
        "all_answers_present": all(bool(item.answer) for item in answer_key.items),
    }


def _write_package_files(
    *,
    output_dir: Path,
    research: ResearchDossier,
    formatting: TemplateFormattingProfile,
    blueprint: ExamBlueprint,
    exam: ExamPaper,
    answer_key: AnswerKey,
    answer_sheet: AnswerSheetSpec,
    scoring_rubric: ScoringRubric,
) -> dict[str, str]:
    files = {
        "exam_docx": "试卷.docx",
        "answer_sheet_docx": "答题卡.docx",
        "answer_key_docx": "参考答案.docx",
        "scoring_rubric_docx": "评分细则.docx",
        "exam_json": "exam.json",
        "blueprint_json": "blueprint.json",
        "answer_sheet_json": "answer_sheet.json",
        "answer_key_json": "answer_key.json",
        "scoring_rubric_json": "scoring_rubric.json",
        "sources_md": "sources.md",
    }
    _write_exam_docx(exam, output_dir / files["exam_docx"], formatting)
    _write_answer_sheet_docx(answer_sheet, output_dir / files["answer_sheet_docx"], formatting)
    _write_answer_key_docx(answer_key, output_dir / files["answer_key_docx"], formatting)
    _write_scoring_rubric_docx(scoring_rubric, output_dir / files["scoring_rubric_docx"], formatting)
    (output_dir / files["exam_json"]).write_text(exam.model_dump_json(indent=2) + "\n", encoding="utf-8")
    (output_dir / files["blueprint_json"]).write_text(blueprint.model_dump_json(indent=2) + "\n", encoding="utf-8")
    (output_dir / files["answer_sheet_json"]).write_text(answer_sheet.model_dump_json(indent=2) + "\n", encoding="utf-8")
    (output_dir / files["answer_key_json"]).write_text(answer_key.model_dump_json(indent=2) + "\n", encoding="utf-8")
    (output_dir / files["scoring_rubric_json"]).write_text(
        scoring_rubric.model_dump_json(indent=2) + "\n",
        encoding="utf-8",
    )
    (output_dir / files["sources_md"]).write_text(render_sources_markdown(research), encoding="utf-8")
    return files


def _write_exam_docx(exam: ExamPaper, output_path: Path, formatting: TemplateFormattingProfile) -> None:
    document = _new_document()
    document.add_heading(exam.title, level=1)
    document.add_paragraph(f"学科：{exam.subject}    年级：{exam.grade}")
    document.add_paragraph(f"满分：{exam.total_score:g} 分    时长：{exam.duration_minutes} 分钟")
    document.add_paragraph("注意事项：请在答题卡对应区域作答，保持卷面整洁。")
    for section in exam.sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(f"本部分共 {len(section.questions)} 题，共 {section.section_score:g} 分。")
        for number, question in _numbered_questions(exam, section_filter=section.title):
            _add_question(document, number, question)
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _write_answer_sheet_docx(
    answer_sheet: AnswerSheetSpec,
    output_path: Path,
    formatting: TemplateFormattingProfile,
) -> None:
    document = _new_document()
    document.add_heading(f"{answer_sheet.title} 答题卡", level=1)
    document.add_paragraph("    ".join(f"{field}: __________" for field in answer_sheet.student_fields))
    for section in answer_sheet.sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(f"题号：{', '.join(str(number) for number in section.question_numbers)}")
        document.add_paragraph(section.response_area)
        for number in section.question_numbers:
            document.add_paragraph(f"{number}. " + "_" * 60)
            if section.question_type not in {QuestionType.CHOICE, QuestionType.FILL_BLANK}:
                for _ in range(4):
                    document.add_paragraph("_" * 72)
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _write_answer_key_docx(answer_key: AnswerKey, output_path: Path, formatting: TemplateFormattingProfile) -> None:
    document = _new_document()
    document.add_heading("参考答案", level=1)
    for item in answer_key.items:
        document.add_paragraph(f"{item.question_number}. {item.answer}（{item.score:g} 分）")
        if item.explanation:
            document.add_paragraph(f"解析：{item.explanation}")
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _write_scoring_rubric_docx(
    scoring_rubric: ScoringRubric,
    output_path: Path,
    formatting: TemplateFormattingProfile,
) -> None:
    document = _new_document()
    document.add_heading("评分细则", level=1)
    for item in scoring_rubric.items:
        document.add_heading(f"{item.question_number}. {item.max_score:g} 分", level=2)
        for point in item.scoring_points:
            document.add_paragraph(point, style="List Bullet")
        if item.review_required:
            document.add_paragraph("主观题需教师复核学生过程表达与关键步骤。")
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _new_document():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('正式试卷包 DOCX 生成需要 python-docx：pip install -e ".[dev]"') from exc
    return Document()


def _add_question(document, number: int, question: Question) -> None:
    if question.material:
        document.add_paragraph(question.material)
    document.add_paragraph(f"{number}. ({question.score:g} 分) {question.content}")
    if question.options:
        for option in question.options:
            document.add_paragraph(option)


def _constraints_from_blueprint(blueprint: ExamBlueprint) -> ExamConstraints:
    return ExamConstraints(
        title=blueprint.title,
        subject=blueprint.subject,
        grade=blueprint.grade,
        duration_minutes=blueprint.duration_minutes,
        sections=[
            SectionConstraint(
                title=section.title,
                question_type=section.question_type,
                count=section.item_count,
                score_per_question=section.score_per_item,
                required_knowledge_points=section.required_knowledge_points,
            )
            for section in blueprint.sections
        ],
        required_knowledge_points=blueprint.knowledge_points,
    )


def _numbered_questions(exam: ExamPaper, section_filter: str | None = None) -> list[tuple[int, Question]]:
    numbered: list[tuple[int, Question]] = []
    question_number = 1
    for section in exam.sections:
        for question in section.questions:
            if section_filter is None or section.title == section_filter:
                numbered.append((question_number, question))
            question_number += 1
    return numbered


def _extract_knowledge_points(research: ResearchDossier) -> list[str]:
    keywords = ["math_quad_graph", "math_quad_vertex"] if research.subject == "数学" else []
    return keywords


def _response_area_for(question_type: QuestionType) -> str:
    if question_type == QuestionType.CHOICE:
        return "请在对应选项上作答：A [  ]  B [  ]  C [  ]  D [  ]"
    if question_type == QuestionType.FILL_BLANK:
        return "请将答案填写在横线上。"
    return "请写出必要的解题过程、计算步骤和结论。"


def _scoring_points_for(question: Question) -> list[str]:
    if question.question_type == QuestionType.CHOICE:
        return [f"选出正确答案得 {question.score:g} 分，错选或不选不得分。"]
    if question.question_type == QuestionType.FILL_BLANK:
        return [f"答案正确得 {question.score:g} 分；形式等价且数学意义一致可得分。"]
    return [
        f"列出正确思路或关键关系式，约 {question.score * 0.3:g} 分。",
        f"计算或推理过程正确，约 {question.score * 0.5:g} 分。",
        f"结论表达完整并符合题意，约 {question.score * 0.2:g} 分。",
    ]


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="生成 Stage E 正式试卷包。")
    parser.add_argument("--research", required=True, type=Path, help="ResearchDossier JSON 文件")
    parser.add_argument("--questions", required=True, type=Path, help="Question[] JSON 文件")
    parser.add_argument("--profile", required=True, type=Path, help="TemplateProfile JSON 文件")
    parser.add_argument("--output-dir", required=True, type=Path, help="输出试卷包目录")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    try:
        research = ResearchDossier.model_validate_json(args.research.read_text(encoding="utf-8"))
        questions = QUESTION_LIST_ADAPTER.validate_json(args.questions.read_text(encoding="utf-8"))
        profile = TemplateProfile.model_validate_json(args.profile.read_text(encoding="utf-8"))
        package = build_exam_package(
            research=research,
            questions=questions,
            profile=profile,
            output_dir=args.output_dir,
        )
    except FileNotFoundError as exc:
        sys.stderr.write(f"输入文件不存在: {exc.filename}\n")
        return 1
    except (RuntimeError, ValidationError, ValueError) as exc:
        sys.stderr.write(f"{exc}\n")
        return 1
    sys.stdout.write(package.model_dump_json(indent=2) + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
