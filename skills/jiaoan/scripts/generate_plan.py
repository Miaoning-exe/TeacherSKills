from __future__ import annotations

import argparse
import sys
from pathlib import Path
from uuid import uuid4

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from pydantic import TypeAdapter, ValidationError

from shared.schemas.lesson import LessonPlan, TeachingStep


LESSON_PLAN_ADAPTER = TypeAdapter(LessonPlan)
ASSET_DIR = Path(__file__).resolve().parent.parent / "assets"
STANDARD_TEMPLATE_PATH = ASSET_DIR / "template_standard.md"
FIVE_E_TEMPLATE_PATH = ASSET_DIR / "template_5e.md"


def parse_beike_report(markdown_text: str) -> dict[str, list[str] | str]:
    section_map = {
        "知识点梳理": "knowledge_points",
        "教学重点与难点": "focus_points",
        "核心素养目标": "competencies",
        "常见误区": "misconceptions",
        "教学策略建议": "strategies",
        "课堂活动建议": "activities",
        "形成性评价建议": "assessments",
    }
    result: dict[str, list[str] | str] = {value: [] for value in section_map.values()}
    result["summary"] = ""

    current_key: str | None = None
    summary_lines: list[str] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("## "):
            heading = line[3:].strip()
            current_key = section_map.get(heading)
            continue
        if current_key is None and line.startswith("- 课标摘要："):
            summary_lines.append(line.removeprefix("- 课标摘要：").strip())
            continue
        if current_key and line.startswith("- "):
            values = result.setdefault(current_key, [])
            if isinstance(values, list):
                values.append(line[2:].strip())
            continue
    result["summary"] = "；".join(summary_lines)
    return result


def generate_lesson_plan(
    *,
    title: str,
    subject: str,
    grade: str,
    template: str,
    duration_minutes: int,
    beike_context: dict[str, list[str] | str] | None = None,
    knowledge_points: list[str] | None = None,
    objectives: list[str] | None = None,
) -> LessonPlan:
    context = beike_context or {}
    merged_knowledge_points = _prefer_user_values(
        explicit_values=knowledge_points,
        fallback_values=_as_list(context.get("knowledge_points")),
    )
    focus_points = _as_list(context.get("focus_points"))
    key_points = [item.removeprefix("教学重点：").strip() for item in focus_points if item.startswith("教学重点：")]
    difficult_points = [item.removeprefix("教学难点：").strip() for item in focus_points if item.startswith("教学难点：")]
    misconceptions = _as_list(context.get("misconceptions"))
    strategies = _as_list(context.get("strategies"))
    activities = _as_list(context.get("activities"))
    assessments = _as_list(context.get("assessments"))
    competencies = _as_list(context.get("competencies"))
    summary = str(context.get("summary", "")).strip()

    final_objectives = _prefer_user_values(
        explicit_values=objectives,
        fallback_values=_build_objectives(
            title=title,
            knowledge_points=merged_knowledge_points,
            competencies=competencies,
        ),
    )
    if not key_points:
        key_points = merged_knowledge_points[: min(3, len(merged_knowledge_points))]
    if not difficult_points:
        difficult_points = misconceptions[: min(3, len(misconceptions))]
    if not difficult_points:
        difficult_points = [f"围绕《{title}》设计梯度问题，避免学生停留在表层记忆。"]

    flow = (
        _generate_five_e_flow(
            title=title,
            duration_minutes=duration_minutes,
            summary=summary,
            knowledge_points=merged_knowledge_points,
            strategies=strategies,
            activities=activities,
            assessments=assessments,
        )
        if template == "5e"
        else _generate_standard_flow(
            title=title,
            duration_minutes=duration_minutes,
            summary=summary,
            knowledge_points=merged_knowledge_points,
            strategies=strategies,
            activities=activities,
            assessments=assessments,
        )
    )

    homework = _build_homework(title=title, knowledge_points=merged_knowledge_points)
    reflection = _build_reflection(template=template, key_points=key_points, difficult_points=difficult_points)
    return LessonPlan(
        id=f"lesson_{uuid4().hex[:12]}",
        title=title,
        subject=subject,
        grade=grade,
        duration_minutes=duration_minutes,
        knowledge_points=merged_knowledge_points,
        objectives=final_objectives,
        key_points=key_points,
        difficult_points=difficult_points,
        teaching_flow=flow,
        homework=homework,
        reflection=reflection,
    )


def lesson_plan_to_json(plan: LessonPlan) -> str:
    return LESSON_PLAN_ADAPTER.dump_json(plan, indent=2).decode("utf-8")


def render_markdown(plan: LessonPlan, *, template: str, assessments: list[str] | None = None) -> str:
    template_path = FIVE_E_TEMPLATE_PATH if template == "5e" else STANDARD_TEMPLATE_PATH
    rendered_flow = _render_teaching_flow(plan.teaching_flow)
    placeholders = {
        "title": plan.title,
        "subject": plan.subject,
        "grade": plan.grade,
        "duration_minutes": str(plan.duration_minutes),
        "knowledge_points": _render_bullets(plan.knowledge_points),
        "objectives": _render_bullets(plan.objectives),
        "key_points": _render_bullets(plan.key_points),
        "difficult_points": _render_bullets(plan.difficult_points),
        "teaching_flow": rendered_flow,
        "assessment": _render_bullets(assessments or ["观察学生表达与任务完成情况，及时调整讲解深度。"]),
        "homework": plan.homework or "请教师补充课后作业。",
        "reflection": plan.reflection or "请课后补充教学反思。",
    }
    return template_path.read_text(encoding="utf-8").format(**placeholders).rstrip() + "\n"


def export_docx(plan: LessonPlan, output_path: Path, *, template: str, assessments: list[str] | None = None) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('DOCX 导出需要安装 python-docx：pip install -e ".[export]"') from exc

    document = Document()
    document.add_heading(plan.title, level=1)
    document.add_paragraph(f"学科：{plan.subject}    年级：{plan.grade}")
    document.add_paragraph(f"课时：{plan.duration_minutes} 分钟    模板：{template}")

    _add_docx_section(document, "教学目标", plan.objectives)
    _add_docx_section(document, "知识点", plan.knowledge_points)
    _add_docx_section(document, "教学重点", plan.key_points)
    _add_docx_section(document, "教学难点", plan.difficult_points)

    document.add_heading("教学过程", level=2)
    for index, step in enumerate(plan.teaching_flow, start=1):
        document.add_paragraph(f"{index}. {step.phase}（{step.duration_minutes} 分钟）")
        document.add_paragraph(f"教学内容：{step.content}")
        document.add_paragraph(f"教师活动：{step.teacher_activity}")
        document.add_paragraph(f"学生活动：{step.student_activity}")
        if step.design_intent:
            document.add_paragraph(f"设计意图：{step.design_intent}")

    if assessments:
        _add_docx_section(document, "课堂评价", assessments)
    if plan.homework:
        document.add_heading("作业设计", level=2)
        document.add_paragraph(plan.homework)
    if plan.reflection:
        document.add_heading("教学反思", level=2)
        document.add_paragraph(plan.reflection)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="根据备课结果或教师输入生成 LessonPlan 教案。")
    parser.add_argument("--title", required=True, help="教案标题")
    parser.add_argument("--subject", required=True, help="学科")
    parser.add_argument("--grade", required=True, help="年级")
    parser.add_argument("--template", choices=["standard", "5e"], default="standard", help="教案模板")
    parser.add_argument("--duration-minutes", type=int, default=45, help="课时长度")
    parser.add_argument("--beike-report", type=Path, default=None, help="备课报告 Markdown")
    parser.add_argument("--knowledge-points", default="", help="逗号分隔的知识点")
    parser.add_argument("--objectives", default="", help="逗号分隔的教学目标")
    parser.add_argument("--output-json", type=Path, default=None, help="输出 LessonPlan JSON")
    parser.add_argument("--output-markdown", type=Path, default=None, help="输出 Markdown 教案")
    parser.add_argument("--output-docx", type=Path, default=None, help="输出 DOCX 教案")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    try:
        beike_context = parse_beike_report(args.beike_report.read_text(encoding="utf-8")) if args.beike_report else None
        knowledge_points = _split_csv(args.knowledge_points)
        objectives = _split_csv(args.objectives)
        plan = generate_lesson_plan(
            title=args.title,
            subject=args.subject,
            grade=args.grade,
            template=args.template,
            duration_minutes=args.duration_minutes,
            beike_context=beike_context,
            knowledge_points=knowledge_points or None,
            objectives=objectives or None,
        )
        assessments = _as_list(beike_context.get("assessments")) if beike_context else []
        markdown = render_markdown(plan, template=args.template, assessments=assessments)
    except FileNotFoundError as exc:
        sys.stderr.write(f"输入文件不存在: {exc.filename}\n")
        return 1
    except (RuntimeError, ValidationError, ValueError) as exc:
        sys.stderr.write(f"{exc}\n")
        return 1

    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(lesson_plan_to_json(plan) + "\n", encoding="utf-8")
    if args.output_markdown:
        args.output_markdown.parent.mkdir(parents=True, exist_ok=True)
        args.output_markdown.write_text(markdown, encoding="utf-8")
    if args.output_docx:
        try:
            export_docx(plan, args.output_docx, template=args.template, assessments=assessments)
        except RuntimeError as exc:
            sys.stderr.write(f"{exc}\n")
            return 1

    if not args.output_markdown and not args.output_json and not args.output_docx:
        sys.stdout.write(markdown)
    return 0


def _prefer_user_values(*, explicit_values: list[str] | None, fallback_values: list[str]) -> list[str]:
    if explicit_values:
        return explicit_values
    return fallback_values or ["请结合教材补充知识点"]


def _build_objectives(*, title: str, knowledge_points: list[str], competencies: list[str]) -> list[str]:
    primary_point = knowledge_points[0] if knowledge_points else title
    objectives = [
        f"理解{primary_point}的核心概念与基本规律。",
        f"能够运用所学内容完成与“{title}”相关的典型任务或练习。",
    ]
    if competencies:
        objectives.append(f"在学习过程中发展{competencies[0]}等核心素养。")
    else:
        objectives.append("在讨论、表达和练习中形成规范的学科思维。")
    return objectives


def _generate_standard_flow(
    *,
    title: str,
    duration_minutes: int,
    summary: str,
    knowledge_points: list[str],
    strategies: list[str],
    activities: list[str],
    assessments: list[str],
) -> list[TeachingStep]:
    durations = _allocate_durations(duration_minutes, [0.15, 0.45, 0.25, 0.15])
    lead_activity = activities[0] if activities else f"围绕“{title}”设置导入问题，引出学习任务。"
    practice_activity = activities[1] if len(activities) > 1 else f"围绕{knowledge_points[0]}设计分层练习。"
    strategy = strategies[0] if strategies else f"结合教材示例逐步讲解《{title}》的关键内容。"
    assessment = assessments[0] if assessments else "通过课堂观察和即时反馈判断学生掌握情况。"
    return [
        TeachingStep(
            phase="导入",
            duration_minutes=durations[0],
            content=f"联系旧知或生活情境，导入《{title}》的学习主题。",
            teacher_activity=lead_activity,
            student_activity="回顾旧知，结合问题表达已有认识。",
            design_intent="激活先备知识，明确学习任务。",
        ),
        TeachingStep(
            phase="新授",
            duration_minutes=durations[1],
            content="聚焦核心概念和关键知识点，完成示范讲解与思路建构。",
            teacher_activity=f"{strategy}；结合课标摘要“{summary or '请教师结合教材补充'}”组织讲解。",
            student_activity=f"记录并归纳 {', '.join(knowledge_points[:3])} 的要点，参与教师引导下的分析。",
            design_intent="完成概念建构，形成清晰的知识框架。",
        ),
        TeachingStep(
            phase="练习",
            duration_minutes=durations[2],
            content="通过分层任务巩固所学，暴露易错点并及时纠偏。",
            teacher_activity=practice_activity,
            student_activity="独立或合作完成练习，并说明解题或表达依据。",
            design_intent="检验学生能否将知识迁移到具体任务中。",
        ),
        TeachingStep(
            phase="小结",
            duration_minutes=durations[3],
            content="梳理本课重点，组织学生总结并布置后续任务。",
            teacher_activity=assessment,
            student_activity="总结收获，提出仍存疑问并记录课后任务。",
            design_intent="帮助学生形成完整认知闭环，为课后巩固做准备。",
        ),
    ]


def _generate_five_e_flow(
    *,
    title: str,
    duration_minutes: int,
    summary: str,
    knowledge_points: list[str],
    strategies: list[str],
    activities: list[str],
    assessments: list[str],
) -> list[TeachingStep]:
    durations = _allocate_durations(duration_minutes, [0.15, 0.2, 0.25, 0.25, 0.15])
    return [
        TeachingStep(
            phase="Engage",
            duration_minutes=durations[0],
            content=f"用问题情境激发对《{title}》的兴趣。",
            teacher_activity=activities[0] if activities else f"用贴近生活的问题引出 {title}。",
            student_activity="观察、猜想并提出初步想法。",
            design_intent="建立问题驱动，激活学生已有经验。",
        ),
        TeachingStep(
            phase="Explore",
            duration_minutes=durations[1],
            content="让学生在任务中自主探索核心现象或规律。",
            teacher_activity=activities[1] if len(activities) > 1 else "组织合作探究，鼓励学生记录发现。",
            student_activity="分组讨论、尝试操作或文本分析，形成初步结论。",
            design_intent="让学生先经历探索，再进入概念提升。",
        ),
        TeachingStep(
            phase="Explain",
            duration_minutes=durations[2],
            content="在教师引导下明确概念、规则和关键方法。",
            teacher_activity=strategies[0] if strategies else f"结合“{summary or title}”提炼关键知识。",
            student_activity=f"表达探究结果，修正并完善对 {', '.join(knowledge_points[:3])} 的理解。",
            design_intent="把经验性认识提升为结构化知识。",
        ),
        TeachingStep(
            phase="Elaborate",
            duration_minutes=durations[3],
            content="通过迁移任务扩展应用场景，提升综合运用能力。",
            teacher_activity=strategies[1] if len(strategies) > 1 else "设计变式任务，引导学生迁移应用。",
            student_activity="在新情境中应用所学并比较不同做法。",
            design_intent="促进知识迁移与深层理解。",
        ),
        TeachingStep(
            phase="Evaluate",
            duration_minutes=durations[4],
            content="通过展示、提问和任务反馈完成形成性评价。",
            teacher_activity=assessments[0] if assessments else "组织学生展示与互评，总结本课收获。",
            student_activity="根据评价标准自评、互评并反思改进。",
            design_intent="及时评估学习效果，形成反馈闭环。",
        ),
    ]


def _build_homework(*, title: str, knowledge_points: list[str]) -> str:
    if knowledge_points:
        return f"完成围绕 {knowledge_points[0]} 的基础巩固题，并选做一题与《{title}》相关的综合应用任务。"
    return f"整理本课《{title}》的知识框架，并完成配套练习。"


def _build_reflection(*, template: str, key_points: list[str], difficult_points: list[str]) -> str:
    return (
        f"课后重点反思 {key_points[0] if key_points else '核心知识'} 的达成情况，关注 {difficult_points[0] if difficult_points else '学生易错点'}。"
        f" 若采用 {template} 模板，可记录各环节时间分配是否合理，以及学生参与度是否达到预期。"
    )


def _allocate_durations(total_minutes: int, ratios: list[float]) -> list[int]:
    raw = [max(1, int(total_minutes * ratio)) for ratio in ratios]
    diff = total_minutes - sum(raw)
    index = 0
    while diff != 0:
        target = index % len(raw)
        if diff > 0:
            raw[target] += 1
            diff -= 1
        elif raw[target] > 1:
            raw[target] -= 1
            diff += 1
        index += 1
    return raw


def _render_bullets(items: list[str]) -> str:
    return "\n".join(f"- {item}" for item in items) if items else "- 请教师补充"


def _render_teaching_flow(steps: list[TeachingStep]) -> str:
    lines: list[str] = []
    for step in steps:
        lines.extend(
            [
                f"### {step.phase}（{step.duration_minutes} 分钟）",
                "",
                f"- 教学内容：{step.content}",
                f"- 教师活动：{step.teacher_activity}",
                f"- 学生活动：{step.student_activity}",
            ]
        )
        if step.design_intent:
            lines.append(f"- 设计意图：{step.design_intent}")
        lines.append("")
    return "\n".join(lines).rstrip()


def _add_docx_section(document, heading: str, items: list[str]) -> None:
    document.add_heading(heading, level=2)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _split_csv(value: str) -> list[str]:
    return [item.strip() for item in value.split(",") if item.strip()]


def _as_list(value: object) -> list[str]:
    return list(value) if isinstance(value, list) else []


if __name__ == "__main__":
    raise SystemExit(main())
