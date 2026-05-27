from __future__ import annotations

import argparse
import sys
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from pydantic import TypeAdapter, ValidationError

from shared.schemas.comment import StudentComment, TeacherObservation
from shared.schemas.knowledge import KnowledgePoint
from shared.schemas.student import KnowledgeMastery, StudentResponse
from shared.schemas.template import TemplateFormattingProfile
from shared.tools.docx_renderer import apply_cjk_formatting


STUDENT_RESPONSE_ADAPTER = TypeAdapter(list[StudentResponse])
KNOWLEDGE_MASTERY_ADAPTER = TypeAdapter(list[KnowledgeMastery])
KNOWLEDGE_POINT_ADAPTER = TypeAdapter(list[KnowledgePoint])
TEACHER_OBSERVATION_ADAPTER = TypeAdapter(list[TeacherObservation])
STUDENT_COMMENT_ADAPTER = TypeAdapter(list[StudentComment])


@dataclass
class StudentCommentProfile:
    student_id: str
    student_name: str | None
    average_ratio: float | None
    average_mastery: float | None
    scored_count: int
    pending_count: int
    strengths: list[str]
    habits: list[str]
    weak_points: list[str]
    improvements: list[str]
    notes: str | None


def load_responses(raw_json: str) -> list[StudentResponse]:
    return STUDENT_RESPONSE_ADAPTER.validate_json(raw_json)


def load_mastery(raw_json: str) -> list[KnowledgeMastery]:
    return KNOWLEDGE_MASTERY_ADAPTER.validate_json(raw_json)


def load_knowledge_points(raw_json: str) -> list[KnowledgePoint]:
    return KNOWLEDGE_POINT_ADAPTER.validate_json(raw_json)


def load_observations(raw_json: str) -> list[TeacherObservation]:
    return TEACHER_OBSERVATION_ADAPTER.validate_json(raw_json)


def generate_student_comments(
    *,
    responses: list[StudentResponse],
    mastery: list[KnowledgeMastery],
    knowledge_points: list[KnowledgePoint],
    observations: list[TeacherObservation] | None = None,
    term: str = "期末",
) -> list[StudentComment]:
    knowledge_lookup = {item.id: item.name for item in knowledge_points}
    observation_lookup = {item.student_id: item for item in observations or []}

    response_group: dict[str, list[StudentResponse]] = defaultdict(list)
    mastery_group: dict[str, list[KnowledgeMastery]] = defaultdict(list)
    for item in responses:
        response_group[item.student_id].append(item)
    for item in mastery:
        mastery_group[item.student_id].append(item)

    student_ids = sorted(set(response_group) | set(mastery_group) | set(observation_lookup))
    comments: list[StudentComment] = []
    for student_id in student_ids:
        profile = build_student_profile(
            student_id=student_id,
            responses=response_group.get(student_id, []),
            mastery=mastery_group.get(student_id, []),
            knowledge_lookup=knowledge_lookup,
            observation=observation_lookup.get(student_id),
        )
        comments.append(render_student_comment(profile=profile, term=term))
    return comments


def build_student_profile(
    *,
    student_id: str,
    responses: list[StudentResponse],
    mastery: list[KnowledgeMastery],
    knowledge_lookup: dict[str, str],
    observation: TeacherObservation | None,
) -> StudentCommentProfile:
    ratios = [
        max(0.0, min((item.score or 0.0) / item.max_score, 1.0))
        for item in responses
        if item.score is not None and item.max_score > 0
    ]
    average_ratio = round(sum(ratios) / len(ratios), 3) if ratios else None
    confident_mastery = [item for item in mastery if _is_confident_mastery(item)]
    average_mastery = (
        round(sum(item.mastery_level for item in confident_mastery) / len(confident_mastery), 3)
        if confident_mastery
        else None
    )
    pending_count = sum(1 for item in responses if item.score is None)

    sorted_mastery = sorted(mastery, key=lambda item: item.mastery_level, reverse=True)
    high_points = [
        knowledge_lookup.get(item.knowledge_point_id, item.knowledge_point_id)
        for item in sorted_mastery
        if item.mastery_level >= 0.75 and _is_confident_mastery(item)
    ][:2]
    weak_points = [
        knowledge_lookup.get(item.knowledge_point_id, item.knowledge_point_id)
        for item in sorted(mastery, key=lambda item: item.mastery_level)
        if item.mastery_level < 0.65 and _is_confident_mastery(item)
    ][:2]

    strengths = _deduplicate((observation.strengths if observation else []) + high_points)
    habits = list(observation.habits) if observation else []
    improvements = _deduplicate((observation.improvements if observation else []) + weak_points)

    return StudentCommentProfile(
        student_id=student_id,
        student_name=observation.student_name if observation else None,
        average_ratio=average_ratio,
        average_mastery=average_mastery,
        scored_count=len(ratios),
        pending_count=pending_count,
        strengths=strengths,
        habits=habits,
        weak_points=weak_points,
        improvements=improvements,
        notes=observation.notes if observation else None,
    )


def render_student_comment(*, profile: StudentCommentProfile, term: str) -> StudentComment:
    variant = _variant_index(profile.student_id)
    student_name = profile.student_name or f"学生{profile.student_id}"
    opener = _OPENER_SENTENCES[variant]
    strength_bridge = _STRENGTH_BRIDGES[variant]
    suggestion_bridge = _SUGGESTION_BRIDGES[variant]
    closing = _CLOSING_SENTENCES[variant]

    overall = _overall_summary(profile)
    habit_text = profile.habits[0] if profile.habits else "能较认真地完成日常学习任务"
    strength_text = _join_items(profile.strengths[:2]) or "课堂学习中表现出稳定的学习基础"
    improvement_text = _join_items(profile.improvements[:2]) or "审题与知识迁移"
    weak_text = _join_items(profile.weak_points[:2])

    sentences = [
        f"{opener}{student_name}在本{term}的学习中{overall}，{habit_text}。",
        f"{strength_bridge}{strength_text}，体现出较好的学习潜力和继续提升的基础。",
        _build_suggestion_sentence(
            suggestion_bridge=suggestion_bridge,
            improvement_text=improvement_text,
            weak_text=weak_text,
            pending_count=profile.pending_count,
        ),
        f"{closing}期待你在下一阶段继续保持自信与投入，把点滴积累转化为更扎实的成长。",
    ]
    if profile.notes:
        sentences.insert(2, f"老师也注意到你{profile.notes}，这说明你已经在积极调整自己的学习状态。")

    comment_text = "".join(sentences)
    return StudentComment(
        student_id=profile.student_id,
        student_name=profile.student_name,
        term=term,
        comment=comment_text,
        highlights=profile.strengths[:3],
        next_steps=profile.improvements[:3],
    )


def comments_to_json(comments: list[StudentComment]) -> str:
    return STUDENT_COMMENT_ADAPTER.dump_json(comments, indent=2).decode("utf-8")


def render_markdown(comments: list[StudentComment], *, term: str) -> str:
    lines = [
        f"# {term}学生评语",
        "",
        f"- 学生人数：{len(comments)}",
        "",
    ]
    for comment in comments:
        display_name = comment.student_name or f"学生 {comment.student_id}"
        lines.extend([f"## {display_name}（{comment.student_id}）", "", comment.comment, ""])
    return "\n".join(lines).rstrip() + "\n"


def export_docx(comments: list[StudentComment], output_path: Path, *, term: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('DOCX 导出需要安装 python-docx：pip install -e ".[export]"') from exc

    document = Document()
    document.add_heading(f"{term}学生评语", level=1)
    document.add_paragraph(f"学生人数：{len(comments)}")
    for comment in comments:
        display_name = comment.student_name or f"学生 {comment.student_id}"
        document.add_heading(f"{display_name}（{comment.student_id}）", level=2)
        document.add_paragraph(comment.comment)

    apply_cjk_formatting(document, TemplateFormattingProfile())
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="根据学情数据和教师观察批量生成学生评语。")
    parser.add_argument("--responses", required=True, type=Path, help="StudentResponse[] JSON 文件")
    parser.add_argument("--mastery", required=True, type=Path, help="KnowledgeMastery[] JSON 文件")
    parser.add_argument("--knowledge-points", required=True, type=Path, help="KnowledgePoint[] JSON 文件")
    parser.add_argument("--observations", type=Path, default=None, help="TeacherObservation[] JSON 文件")
    parser.add_argument("--term", default="期末", help="评语场景")
    parser.add_argument("--output-json", type=Path, default=None, help="输出 StudentComment[] JSON")
    parser.add_argument("--output-markdown", type=Path, default=None, help="输出 Markdown 评语")
    parser.add_argument("--output-docx", type=Path, default=None, help="输出 DOCX 评语")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    try:
        responses = load_responses(args.responses.read_text(encoding="utf-8"))
        mastery = load_mastery(args.mastery.read_text(encoding="utf-8"))
        knowledge_points = load_knowledge_points(args.knowledge_points.read_text(encoding="utf-8"))
        observations = load_observations(args.observations.read_text(encoding="utf-8")) if args.observations else None
        comments = generate_student_comments(
            responses=responses,
            mastery=mastery,
            knowledge_points=knowledge_points,
            observations=observations,
            term=args.term,
        )
        markdown = render_markdown(comments, term=args.term)
    except FileNotFoundError as exc:
        sys.stderr.write(f"输入文件不存在: {exc.filename}\n")
        return 1
    except ValidationError as exc:
        sys.stderr.write(f"{exc}\n")
        return 1

    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(comments_to_json(comments) + "\n", encoding="utf-8")
    if args.output_markdown:
        args.output_markdown.parent.mkdir(parents=True, exist_ok=True)
        args.output_markdown.write_text(markdown, encoding="utf-8")
    if args.output_docx:
        try:
            export_docx(comments, args.output_docx, term=args.term)
        except RuntimeError as exc:
            sys.stderr.write(f"{exc}\n")
            return 1

    if not args.output_json and not args.output_markdown and not args.output_docx:
        sys.stdout.write(markdown)
    return 0


def _overall_summary(profile: StudentCommentProfile) -> str:
    if profile.average_ratio is None:
        return "保持着稳定的学习态度"
    if profile.scored_count <= 1 and profile.pending_count >= 1:
        return "展现出一定的学习潜力"
    if profile.pending_count >= max(2, profile.scored_count):
        return "阶段表现有亮点，也还有继续观察和提升的空间"
    if profile.average_mastery is not None and profile.average_mastery < 0.6:
        return "正在逐步夯实基础"
    if profile.average_ratio >= 0.85:
        return "整体表现稳健，学习成效较为突出"
    if profile.average_ratio >= 0.7:
        return "整体表现较好，基础比较扎实"
    if profile.average_ratio >= 0.55:
        return "呈现出持续进步的态势"
    return "仍处在夯实基础、逐步追赶的阶段"


def _build_suggestion_sentence(
    *,
    suggestion_bridge: str,
    improvement_text: str,
    weak_text: str,
    pending_count: int,
) -> str:
    pending_note = ""
    if pending_count > 0:
        pending_note = "部分综合表现还有待结合后续批改继续观察，"
    if weak_text:
        return f"{suggestion_bridge}{pending_note}如果能在{improvement_text}方面再主动一些，尤其关注{weak_text}的巩固，进步会更加明显。"
    return f"{suggestion_bridge}{pending_note}如果能在{improvement_text}方面继续坚持，你的学习状态会更加稳定。"


def _join_items(items: list[str]) -> str:
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return "、".join(items)


def _variant_index(student_id: str) -> int:
    return sum(ord(char) for char in student_id) % len(_OPENER_SENTENCES)


def _deduplicate(items: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for item in items:
        cleaned = item.strip()
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            ordered.append(cleaned)
    return ordered


def _is_confident_mastery(item: KnowledgeMastery) -> bool:
    confidence = item.confidence if item.confidence is not None else 0.0
    return confidence >= 0.4


_OPENER_SENTENCES = [
    "你在老师眼中一直是一个愿意投入学习的孩子，",
    "这一阶段的学习中，",
    "回顾本学期的表现，",
]

_STRENGTH_BRIDGES = [
    "从课堂表现和学情结果来看，你在",
    "从最近的学习状态来看，你在",
    "从平时作答和课堂反馈来看，你在",
]

_SUGGESTION_BRIDGES = [
    "接下来，",
    "后续学习中，",
    "如果放眼下一阶段，",
]

_CLOSING_SENTENCES = [
    "老师相信你有能力把优势继续放大，",
    "只要保持现在的节奏，",
    "愿你带着这份积累继续向前，",
]


if __name__ == "__main__":
    raise SystemExit(main())
