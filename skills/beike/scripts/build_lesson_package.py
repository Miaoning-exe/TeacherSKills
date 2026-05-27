from __future__ import annotations

import argparse
import sys
import uuid
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from pydantic import ValidationError

from shared.schemas.lesson import LessonContext, LessonPackage, LessonPlan
from shared.schemas.research import CredibilityLevel, ResearchDossier, SourceEvidence, SourceType
from shared.schemas.template import TemplateFormattingProfile
from shared.tools.docx_renderer import apply_cjk_formatting
from shared.tools.sources import render_sources_markdown
from skills.beike.scripts.analyze_curriculum import (
    DEFAULT_BLOOM_PATH,
    DEFAULT_CURRICULUM_PATH,
    CurriculumAnalysis,
    analyze_curriculum,
    generate_analysis_report,
    load_bloom_descriptions,
    load_curriculum_entries,
)
from skills.jiaoan.scripts.generate_plan import (
    export_docx,
    generate_lesson_plan,
    lesson_context_to_beike_context,
)


def build_lesson_package(
    *,
    research: ResearchDossier,
    output_dir: Path,
    curriculum_file: Path = DEFAULT_CURRICULUM_PATH,
    bloom_file: Path = DEFAULT_BLOOM_PATH,
    template: str = "standard",
    duration_minutes: int = 45,
    formatting: TemplateFormattingProfile | None = None,
) -> LessonPackage:
    formatting = formatting or TemplateFormattingProfile()
    curriculum_text = curriculum_file.read_text(encoding="utf-8")
    bloom_text = bloom_file.read_text(encoding="utf-8")
    entries = load_curriculum_entries(curriculum_text)
    keywords = _keywords_from_research(research)
    analysis = analyze_curriculum(
        subject=research.subject,
        grade=research.grade,
        topic=research.topic,
        keywords=keywords,
        entries=entries,
    )
    report = generate_analysis_report(
        analysis=analysis,
        bloom_descriptions=load_bloom_descriptions(bloom_text),
        research=research,
    )
    context = build_lesson_context(
        research=research,
        analysis=analysis,
        duration_minutes=duration_minutes,
    )
    plan = generate_lesson_plan(
        title=context.title,
        subject=context.subject,
        grade=context.grade,
        template=template,
        duration_minutes=context.duration_minutes,
        beike_context=lesson_context_to_beike_context(context),
    )

    output_dir.mkdir(parents=True, exist_ok=True)
    files = {
        "beike_analysis_docx": "备课分析.docx",
        "teaching_design_docx": "教学设计.docx",
        "activity_sheet_docx": "课堂活动单.docx",
        "practice_docx": "配套练习.docx",
        "lesson_context_json": "lesson_context.json",
        "lesson_plan_json": "lesson_plan.json",
        "package_json": "package.json",
        "sources_md": "sources.md",
    }
    _write_beike_analysis_docx(context, report, output_dir / files["beike_analysis_docx"], formatting)
    export_docx(plan, output_dir / files["teaching_design_docx"], template=template, assessments=context.assessment_suggestions, formatting=formatting)
    _write_activity_sheet_docx(context, output_dir / files["activity_sheet_docx"], formatting)
    _write_practice_docx(context, plan, output_dir / files["practice_docx"], formatting)
    (output_dir / files["lesson_context_json"]).write_text(context.model_dump_json(indent=2) + "\n", encoding="utf-8")
    (output_dir / files["lesson_plan_json"]).write_text(plan.model_dump_json(indent=2) + "\n", encoding="utf-8")
    (output_dir / files["sources_md"]).write_text(render_sources_markdown(research), encoding="utf-8")

    checks = {
        "context_has_sources": bool(context.source_ids),
        "context_has_review_notes": bool(context.teacher_review_notes),
        "plan_uses_context_title": plan.title == context.title,
        "plan_has_teaching_flow": bool(plan.teaching_flow),
    }
    checks.update({f"file_exists:{key}": (output_dir / filename).exists() for key, filename in files.items() if key != "package_json"})
    package = LessonPackage(
        id=f"lesson_package_{uuid.uuid4().hex[:12]}",
        lesson_context_id=context.id,
        lesson_plan_id=plan.id,
        created_at=datetime.now().isoformat(timespec="seconds"),
        output_dir=str(output_dir),
        files=files,
        checks=checks,
        source_ids=context.source_ids,
    )
    (output_dir / files["package_json"]).write_text(package.model_dump_json(indent=2) + "\n", encoding="utf-8")
    return package


def build_lesson_context(
    *,
    research: ResearchDossier,
    analysis: CurriculumAnalysis,
    duration_minutes: int = 45,
) -> LessonContext:
    matched_entries = analysis.matched_entries
    return LessonContext(
        id=f"lesson_context_{uuid.uuid4().hex[:12]}",
        title=research.topic,
        subject=research.subject,
        grade=research.grade,
        topic=research.topic,
        duration_minutes=duration_minutes,
        curriculum_alignment=[entry.standard_summary for entry in matched_entries],
        knowledge_points=_deduplicate(item for entry in matched_entries for item in entry.knowledge_points),
        core_competencies=_deduplicate(item for entry in matched_entries for item in entry.core_competencies),
        key_points=analysis.key_points or research.key_findings[:2],
        difficult_points=analysis.difficult_points or research.key_findings[2:4],
        misconceptions=_deduplicate(item for entry in matched_entries for item in entry.misconceptions),
        teaching_strategies=_deduplicate(item for entry in matched_entries for item in entry.teaching_strategies)
        + research.agent_inferences[:2],
        activity_suggestions=_deduplicate(item for entry in matched_entries for item in entry.activity_suggestions),
        assessment_suggestions=_deduplicate(item for entry in matched_entries for item in entry.assessment_suggestions),
        teacher_review_notes=research.teacher_review_notes
        or [source.review_note for source in research.sources if source.review_note],
        source_ids=[source.id for source in research.sources],
        local_fallback=research.local_fallback,
    )


def build_local_fallback_research(*, subject: str, grade: str, topic: str, keywords: list[str]) -> ResearchDossier:
    return ResearchDossier(
        id=f"research_local_beike_{uuid.uuid4().hex[:8]}",
        task_type="备课",
        subject=subject,
        grade=grade,
        topic=topic,
        created_at=datetime.now().date().isoformat(),
        query_summary=f"本地参考降级：{subject}{grade}{topic}，关键词：{', '.join(keywords) or '无'}。",
        local_fallback=True,
        sources=[
            SourceEvidence(
                id="src_local_curriculum_reference",
                title="TeacherSkills 本地课标摘录",
                source_type=SourceType.LOCAL_REFERENCE,
                summary="使用本仓库本地课标摘录与 Bloom 分类参考生成备课草稿。",
                publisher="TeacherSkills",
                retrieved_at=datetime.now().date().isoformat(),
                credibility=CredibilityLevel.MEDIUM,
                citation_locations=["备课分析", "教学设计"],
                review_note="该来源不是实时检索资料，请教师结合本校教材与最新课标复核。",
            )
        ],
        key_findings=[],
        agent_inferences=[],
        teacher_review_notes=["本产物为本地模板草稿，需要教师补充真实教材页码、例题编号和班级学情。"],
    )


def _write_beike_analysis_docx(
    context: LessonContext,
    report: str,
    output_path: Path,
    formatting: TemplateFormattingProfile,
) -> None:
    document = _new_document()
    document.add_heading("备课分析", level=1)
    document.add_paragraph(f"学科：{context.subject}    年级：{context.grade}    主题：{context.topic}")
    for line in report.splitlines():
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|") or not line.strip():
            continue
        else:
            document.add_paragraph(line)
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _write_activity_sheet_docx(
    context: LessonContext,
    output_path: Path,
    formatting: TemplateFormattingProfile,
) -> None:
    document = _new_document()
    document.add_heading(f"{context.title} 课堂活动单", level=1)
    document.add_paragraph("姓名：__________    小组：__________    日期：__________")
    document.add_heading("学习目标", level=2)
    for point in context.knowledge_points[:3]:
        document.add_paragraph(f"理解并应用：{point}", style="List Bullet")
    document.add_heading("活动任务", level=2)
    activities = context.activity_suggestions or ["围绕本课主题完成观察、讨论和表达任务。"]
    for index, activity in enumerate(activities, start=1):
        document.add_paragraph(f"任务 {index}：{activity}")
        document.add_paragraph("我的记录：")
        for _ in range(3):
            document.add_paragraph("_" * 64)
    document.add_heading("课堂自评", level=2)
    for item in context.assessment_suggestions or ["我能说出本课核心知识并完成基础练习。"]:
        document.add_paragraph(f"[  ] {item}")
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _write_practice_docx(
    context: LessonContext,
    plan: LessonPlan,
    output_path: Path,
    formatting: TemplateFormattingProfile,
) -> None:
    document = _new_document()
    document.add_heading(f"{context.title} 配套练习", level=1)
    document.add_paragraph(f"建议用时：15 分钟    对应课时：{plan.duration_minutes} 分钟")
    points = context.knowledge_points or [context.topic]
    document.add_heading("基础巩固", level=2)
    for index, point in enumerate(points[:3], start=1):
        document.add_paragraph(f"{index}. 请用自己的话说明“{point}”的含义，并举一个例子。")
    document.add_heading("迁移应用", level=2)
    document.add_paragraph(f"4. 结合《{context.title}》中的核心方法，解释一个教材或生活中的相关情境。")
    document.add_heading("反思提升", level=2)
    for index, point in enumerate(context.difficult_points[:2], start=5):
        document.add_paragraph(f"{index}. 针对易错点“{point}”，写出你容易出错的原因和改进办法。")
    apply_cjk_formatting(document, formatting)
    document.save(output_path)


def _new_document():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError('备课文档包 DOCX 生成需要 python-docx：pip install -e ".[dev]"') from exc
    return Document()


def _keywords_from_research(research: ResearchDossier) -> list[str]:
    words = [research.topic]
    words.extend(finding for finding in research.key_findings if len(finding) <= 20)
    return words


def _deduplicate(items) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for item in items:
        if item and item not in seen:
            seen.add(item)
            ordered.append(item)
    return ordered


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="生成 Stage E 正式备课包。")
    parser.add_argument("--research", type=Path, default=None, help="ResearchDossier JSON 文件")
    parser.add_argument("--subject", default="", help="无资料包时使用的学科")
    parser.add_argument("--grade", default="", help="无资料包时使用的年级")
    parser.add_argument("--topic", default="", help="无资料包时使用的主题")
    parser.add_argument("--keywords", default="", help="无资料包时使用的逗号分隔关键词")
    parser.add_argument("--template", choices=["standard", "5e"], default="standard", help="教案模板")
    parser.add_argument("--duration-minutes", type=int, default=45, help="课时长度")
    parser.add_argument("--curriculum-file", type=Path, default=DEFAULT_CURRICULUM_PATH, help="课标参考文件")
    parser.add_argument("--bloom-file", type=Path, default=DEFAULT_BLOOM_PATH, help="Bloom 分类参考文件")
    parser.add_argument("--output-dir", required=True, type=Path, help="输出备课包目录")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    try:
        keywords = [item.strip() for item in args.keywords.split(",") if item.strip()]
        if args.research:
            research = ResearchDossier.model_validate_json(args.research.read_text(encoding="utf-8"))
        else:
            if not args.subject or not args.grade or not args.topic:
                raise ValueError("未提供 --research 时，--subject、--grade、--topic 均为必填")
            research = build_local_fallback_research(
                subject=args.subject,
                grade=args.grade,
                topic=args.topic,
                keywords=keywords,
            )
        package = build_lesson_package(
            research=research,
            output_dir=args.output_dir,
            curriculum_file=args.curriculum_file,
            bloom_file=args.bloom_file,
            template=args.template,
            duration_minutes=args.duration_minutes,
        )
    except FileNotFoundError as exc:
        sys.stderr.write(f"输入文件不存在: {exc.filename}\n")
        return 1
    except (RuntimeError, ValidationError, ValueError) as exc:
        sys.stderr.write(f"{exc}\n")
        return 1
    sys.stdout.write(package.model_dump_json(indent=2) + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
